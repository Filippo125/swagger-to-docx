
import json
from docx import Document
from swagger_parser import SwaggerParser


def build_api_section(swagger: SwaggerParser, doc: Document):
	doc.add_heading('RestApi', level=2)
	for path, path_api in swagger.get_path().items():
		for method, m_api in path_api.items():
			# if available use summary as api title, if not use operationId, if not use empty string
			summary = m_api.get("summary", m_api.get("operationId", ""))
			doc.add_heading(summary, level=3)
			doc.add_paragraph(m_api.get("description", ""))
			p = doc.add_paragraph("")
			if "x-scope" in m_api:
				p.add_run('Scope: \n').bold = True
				for scope in m_api["x-scope"]:
					p.add_run("- %s\n" % scope)
			p.add_run('Path: ').bold = True
			p.add_run('%s\n' %(path))
			p.add_run('Method: ').bold = True
			p.add_run("%s\n" % (method.upper()))
			body_ref = swagger.get_body_ref(m_api)
			p.add_run('Body: ').bold = True
			if body_ref:
				# doc.add_paragraph("Body: \n" )#, style='List Bullet')
				table = doc.add_table(rows=1, cols=1)
				table.rows[0].cells[0].text = json.dumps(swagger.build_example_object(swagger.get_ref(body_ref)), indent=2)
			else:
				p.add_run("No body content is required")
			doc.add_paragraph("").add_run("Responses:").bold = True
			table = doc.add_table(rows=1, cols=3)
			table.style = 'Light List Accent 1'
			hdr_cells = table.rows[0].cells
			hdr_cells[0].text = 'Response Code'
			hdr_cells[1].text = 'Description'
			hdr_cells[2].text = 'Content'
			for code, resp in m_api["responses"].items():
				row_cells = table.add_row().cells
				row_cells[0].text = code
				row_cells[1].text = resp.get("description", "")
				if "schema" in resp:
					if "$ref" in resp["schema"]:
						row_cells[2].text = json.dumps(swagger.build_example_object(swagger.get_ref(resp["schema"]["$ref"])), indent=2)
					else:
						row_cells[2].text = json.dumps(swagger.build_example_object(resp["schema"]), indent=2)
			p = doc.add_paragraph("")
			p.add_run('Security:\n').bold = True
			if m_api.get("security"):
				perms = swagger.get_permission_list(m_api)
				if perms:
					p.add_run("The following permission are required to complete the operation:")
					for perm in perms:
						doc.add_paragraph(perm, style='List Bullet')
				else:
					p.add_run("Only authentication is required")
			else:
				p.add_run("There are not required permission or autentication")


def build_model_section(swagger: SwaggerParser, doc: Document):
	doc.add_heading('Models', level=2)
	doc.add_paragraph("In this section all models used are described")
	definitions = swagger.get_definitions()
	for model, value in definitions.items():
		doc.add_heading(model, level=3)
		doc.add_paragraph(value.get("description",""))
		table = doc.add_table(rows=1, cols=3)
		table.style = 'Light List Accent 1'
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = 'Field'
		hdr_cells[1].text = 'Type'
		hdr_cells[2].text = 'Description'
		# basic object
		properties = swagger.get_model_triplette_list(model_value=value)
		for prop in properties:
			row_cells = table.add_row().cells
			row_cells[0].text = prop["name"]
			row_cells[1].text = prop["type"]
			row_cells[2].text = prop["desc"]



def build_title_section(swagger: SwaggerParser, doc: Document):
	doc.add_heading(swagger.get_info()["title"], level=0)
	doc.add_paragraph(swagger.get_info().get("description",""))

